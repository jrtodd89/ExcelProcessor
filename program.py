import os
import pandas as pd
from docx import Document

currentDirectory = os.path.dirname(os.path.abspath(__file__))
dataPath = os.path.join(currentDirectory+'/data/')
resultsPath = os.path.join(currentDirectory+'/results/')

print(dataPath)
print(resultsPath)


allStudents = pd.read_csv("data/students.csv")
allStudents.rename(columns = {'Hours Remaining': 'Hours'}, inplace = True)
students = allStudents[allStudents['Hours'] < 100]

for student in students.itertuples(index=False, name='Student'):
	name = student.Name
	address = student.Address
	grade = student.Grade
	info = student.Info
	hours = student.Hours

	document = Document()
	heading = document.add_heading('Test Heading')

	nameString = f'The Quick Brown {name} Jumps Over The Lazy Dog.'
	paragraph = document.add_paragraph(nameString)


	print(nameString)
	document.save(f'{resultsPath}{name}.docx')